#!/usr/bin/env python
# -*- coding: utf-8 -*-

"""
Medical AI Tools - PDF/DOC Reader Tool
-----------------------------------
Tool đọc và parse PDF/DOC files cho RAG system.
"""

import os
import json
import pickle
from typing import Dict, Any, List, Optional
from pathlib import Path
import hashlib

from medical_ai_agents.tools.base_tools import BaseTool

class PDFDocReaderTool(BaseTool):
    """Tool đọc và parse PDF/DOC files."""
    
    def __init__(self, storage_path: str, chunk_size: int = 500, overlap: int = 50, **kwargs):
        """Initialize PDF/DOC reader tool."""
        super().__init__(
            name="pdf_doc_reader",
            description="Đọc và parse PDF/DOC files thành chunks cho RAG system."
        )
        
        self.storage_path = Path(storage_path)
        self.storage_path.mkdir(exist_ok=True)
        self.chunk_size = chunk_size
        self.overlap = overlap
        
        # Storage paths
        self.chunks_path = self.storage_path / "chunks"
        self.chunks_path.mkdir(exist_ok=True)
        self.metadata_path = self.storage_path / "metadata.json"
        
        # Track processed files
        self.processed_files = self._load_processed_files()
    
    def _load_processed_files(self) -> Dict[str, Any]:
        """Load metadata about processed files."""
        if self.metadata_path.exists():
            try:
                with open(self.metadata_path, 'r', encoding='utf-8') as f:
                    return json.load(f)
            except:
                return {}
        return {}
    
    def _save_processed_files(self):
        """Save metadata about processed files."""
        with open(self.metadata_path, 'w', encoding='utf-8') as f:
            json.dump(self.processed_files, f, ensure_ascii=False, indent=2)
    
    def _get_file_hash(self, file_path: str) -> str:
        """Get hash of file for deduplication."""
        hasher = hashlib.md5()
        with open(file_path, 'rb') as f:
            buf = f.read(65536)
            while len(buf) > 0:
                hasher.update(buf)
                buf = f.read(65536)
        return hasher.hexdigest()
    
    def _run(self, file_paths: List[str]) -> Dict[str, Any]:
        """Process PDF/DOC files."""
        try:
            documents_processed = []
            total_chunks = 0
            errors = []
            
            for file_path in file_paths:
                if not os.path.exists(file_path):
                    errors.append(f"File not found: {file_path}")
                    continue
                
                # Check if already processed
                file_hash = self._get_file_hash(file_path)
                if file_hash in self.processed_files:
                    self.logger.info(f"File already processed: {file_path}")
                    documents_processed.append(os.path.basename(file_path))
                    continue
                
                # Process based on file type
                file_ext = Path(file_path).suffix.lower()
                
                if file_ext == '.pdf':
                    chunks = self._process_pdf(file_path)
                elif file_ext in ['.doc', '.docx']:
                    chunks = self._process_doc(file_path)
                elif file_ext == '.txt':
                    chunks = self._process_text(file_path)
                else:
                    errors.append(f"Unsupported file type: {file_ext}")
                    continue
                
                # Save chunks
                if chunks:
                    chunk_file = self.chunks_path / f"{file_hash}.pkl"
                    with open(chunk_file, 'wb') as f:
                        pickle.dump(chunks, f)
                    
                    # Update metadata
                    self.processed_files[file_hash] = {
                        "file_name": os.path.basename(file_path),
                        "file_path": file_path,
                        "chunks_count": len(chunks),
                        "file_type": file_ext,
                        "processed_at": str(Path(file_path).stat().st_mtime)
                    }
                    
                    documents_processed.append(os.path.basename(file_path))
                    total_chunks += len(chunks)
                    
                    self.logger.info(f"Processed {file_path}: {len(chunks)} chunks")
            
            # Save metadata
            self._save_processed_files()
            
            return {
                "success": True,
                "documents_processed": documents_processed,
                "total_chunks": total_chunks,
                "errors": errors if errors else None
            }
            
        except Exception as e:
            import traceback
            return {
                "success": False,
                "error": str(e),
                "traceback": traceback.format_exc()
            }
    
    def _process_pdf(self, file_path: str) -> List[Dict[str, Any]]:
        """Process PDF file."""
        chunks = []
        
        try:
            # Try PyMuPDF first (faster)
            try:
                import fitz  # PyMuPDF
                doc = fitz.open(file_path)
                
                full_text = ""
                page_texts = []
                
                for page_num in range(len(doc)):
                    page = doc[page_num]
                    text = page.get_text()
                    page_texts.append((page_num + 1, text))
                    full_text += f"\n--- Page {page_num + 1} ---\n{text}"
                
                doc.close()
                
            except ImportError:
                # Fallback to PyPDF2
                try:
                    import PyPDF2
                    with open(file_path, 'rb') as file:
                        pdf_reader = PyPDF2.PdfReader(file)
                        
                        full_text = ""
                        page_texts = []
                        
                        for page_num, page in enumerate(pdf_reader.pages):
                            text = page.extract_text()
                            page_texts.append((page_num + 1, text))
                            full_text += f"\n--- Page {page_num + 1} ---\n{text}"
                            
                except ImportError:
                    self.logger.error("No PDF reader available. Install PyMuPDF or PyPDF2")
                    return []
            
            # Create chunks with page info
            for page_num, page_text in page_texts:
                page_chunks = self._split_text_into_chunks(page_text)
                
                for i, chunk_text in enumerate(page_chunks):
                    if len(chunk_text.strip()) > 50:  # Min chunk size
                        chunk = {
                            "content": chunk_text.strip(),
                            "source": os.path.basename(file_path),
                            "page": page_num,
                            "chunk_id": f"p{page_num}_c{i}",
                            "metadata": {
                                "file_type": "pdf",
                                "full_path": file_path
                            }
                        }
                        chunks.append(chunk)
            
            return chunks
            
        except Exception as e:
            self.logger.error(f"Error processing PDF {file_path}: {str(e)}")
            return []
    
    def _process_doc(self, file_path: str) -> List[Dict[str, Any]]:
        """Process DOC/DOCX file."""
        chunks = []
        
        try:
            import docx
            doc = docx.Document(file_path)
            
            # Extract all text
            full_text = ""
            paragraphs = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
                    full_text += para.text + "\n\n"
            
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells)
                    if row_text.strip():
                        full_text += row_text + "\n"
            
            # Create chunks
            text_chunks = self._split_text_into_chunks(full_text)
            
            for i, chunk_text in enumerate(text_chunks):
                if len(chunk_text.strip()) > 50:
                    chunk = {
                        "content": chunk_text.strip(),
                        "source": os.path.basename(file_path),
                        "page": 0,  # DOC doesn't have pages
                        "chunk_id": f"c{i}",
                        "metadata": {
                            "file_type": "docx",
                            "full_path": file_path
                        }
                    }
                    chunks.append(chunk)
            
            return chunks
            
        except ImportError:
            self.logger.error("python-docx not installed. Cannot process DOC files.")
            return []
        except Exception as e:
            self.logger.error(f"Error processing DOC {file_path}: {str(e)}")
            return []
    
    def _process_text(self, file_path: str) -> List[Dict[str, Any]]:
        """Process plain text file."""
        chunks = []
        
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # Create chunks
            text_chunks = self._split_text_into_chunks(content)
            
            for i, chunk_text in enumerate(text_chunks):
                if len(chunk_text.strip()) > 50:
                    chunk = {
                        "content": chunk_text.strip(),
                        "source": os.path.basename(file_path),
                        "page": 0,
                        "chunk_id": f"c{i}",
                        "metadata": {
                            "file_type": "txt",
                            "full_path": file_path
                        }
                    }
                    chunks.append(chunk)
            
            return chunks
            
        except Exception as e:
            self.logger.error(f"Error processing text file {file_path}: {str(e)}")
            return []
    
    def _split_text_into_chunks(self, text: str) -> List[str]:
        """Split text into overlapping chunks."""
        if len(text) <= self.chunk_size:
            return [text]
        
        chunks = []
        sentences = self._split_into_sentences(text)
        
        current_chunk = []
        current_length = 0
        
        for sentence in sentences:
            sentence_length = len(sentence)
            
            if current_length + sentence_length <= self.chunk_size:
                current_chunk.append(sentence)
                current_length += sentence_length
            else:
                # Save current chunk
                if current_chunk:
                    chunks.append(" ".join(current_chunk))
                
                # Start new chunk with overlap
                if self.overlap > 0 and len(current_chunk) > 1:
                    # Keep last few sentences for overlap
                    overlap_sentences = current_chunk[-(self.overlap // 100):]
                    current_chunk = overlap_sentences + [sentence]
                    current_length = sum(len(s) for s in current_chunk)
                else:
                    current_chunk = [sentence]
                    current_length = sentence_length
        
        # Add last chunk
        if current_chunk:
            chunks.append(" ".join(current_chunk))
        
        return chunks
    
    def _split_into_sentences(self, text: str) -> List[str]:
        """Simple sentence splitter."""
        # Simple approach - split by punctuation
        import re
        sentences = re.split(r'(?<=[.!?])\s+', text)
        
        # Filter empty sentences
        sentences = [s.strip() for s in sentences if s.strip()]
        
        # Handle very long sentences
        final_sentences = []
        for sentence in sentences:
            if len(sentence) > self.chunk_size:
                # Split by newlines or commas
                parts = re.split(r'\n|,\s*', sentence)
                final_sentences.extend(parts)
            else:
                final_sentences.append(sentence)
        
        return final_sentences
    
    def get_parameters_schema(self) -> Dict[str, Any]:
        """Return JSON schema for the tool parameters."""
        return {
            "file_paths": {
                "type": "array",
                "description": "List of file paths to process",
                "items": {"type": "string"}
            }
        }